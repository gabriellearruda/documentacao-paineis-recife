"""
Gera relatorio_gestor.docx e relatorio_tecnico.docx a partir dos .md.
Usa python-docx. Insere fluxograma ao final de cada secao de pagina.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS    = "C:/Users/gabri/painel-shiny-recife/docs"
IMAGENS = f"{DOCS}/imagens"

AZUL     = RGBColor(0x12, 0x1E, 0x87)
VERMELHO = RGBColor(0xFF, 0x50, 0x54)
CINZA    = RGBColor(0x6b, 0x72, 0x80)
BRANCO   = RGBColor(0xFF, 0xFF, 0xFF)

FLUXO_MAP = {
    "mapa":          f"{IMAGENS}/fluxo_p1_mapa.png",
    "sinan viol":    f"{IMAGENS}/fluxo_p2_sinan_viol.png",
    "intoxica":      f"{IMAGENS}/fluxo_p3_iexo.png",
    "hospitaliza":   f"{IMAGENS}/fluxo_p4_sih.png",
    "mortalidade":   f"{IMAGENS}/fluxo_p5_sim.png",
    "linkage":       f"{IMAGENS}/fluxo_p6_linkage.png",
    "linha da vida": f"{IMAGENS}/fluxo_p7_linha_vida.png",
    "subnotifica":   f"{IMAGENS}/fluxo_p8_subnotificacoes.png",
}


# ── Helpers XML ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horiz_rule(doc, color="E5E7EB"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── Estilos ────────────────────────────────────────────────────────────────────
def setup_styles(doc):
    from docx.oxml.ns import nsmap
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for name, size, color, space_before, space_after in [
        ("Heading 1", 16, AZUL,   Pt(18), Pt(6)),
        ("Heading 2", 12, AZUL,   Pt(14), Pt(4)),
        ("Heading 3", 10, RGBColor(0x1f, 0x29, 0x37), Pt(10), Pt(2)),
    ]:
        s = doc.styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = space_before
        s.paragraph_format.space_after = space_after
        s.paragraph_format.keep_with_next = True


# ── Capa ───────────────────────────────────────────────────────────────────────
def add_capa(doc, titulo, subtitulo):
    # Fundo azul simulado via paragrafo com shading seria complexo;
    # fazemos uma pagina de titulo limpa.
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = AZUL

    # Linha vermelha
    add_horiz_rule(doc, "FF5054")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for linha in subtitulo.split("\n"):
        run2 = p2.add_run(linha + "\n")
        run2.font.name = "Arial"
        run2.font.size = Pt(12)
        run2.font.color.rgb = CINZA

    page_break(doc)


# ── Tabela Markdown ────────────────────────────────────────────────────────────
def add_table(doc, linhas_tabela):
    linhas = [l for l in linhas_tabela
              if l.strip() and not re.match(r"^\|[-| :]+\|$", l.strip())]
    if not linhas:
        return

    rows = []
    for l in linhas:
        cells = [re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1",
                 re.sub(r"`(.+?)`", r"\1", c.strip()))
                 for c in l.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            txt = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = BRANCO
                set_cell_bg(cell, "121E87")
            else:
                bg = "F9FAFB" if i % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)

    doc.add_paragraph()


# ── Fluxograma ────────────────────────────────────────────────────────────────
def fluxo_para_titulo(titulo):
    t = titulo.lower()
    for chave, caminho in FLUXO_MAP.items():
        if chave in t:
            return caminho
    return None


def insert_fluxo(doc, titulo_secao):
    img = fluxo_para_titulo(titulo_secao)
    if not img or not os.path.exists(img):
        return
    p_label = doc.add_paragraph()
    run = p_label.add_run("Fluxo de dados desta página:")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = CINZA
    p_label.paragraph_format.space_before = Pt(6)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img, width=Inches(6.2))
    doc.add_paragraph()


# ── Parser Markdown ────────────────────────────────────────────────────────────
def render_md(doc, md_text, inserir_fluxo=False):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_buf = []
    secao_atual = None

    def flush_table():
        if table_buf:
            add_table(doc, list(table_buf))
            table_buf.clear()

    def flush_fluxo():
        if inserir_fluxo and secao_atual:
            insert_fluxo(doc, secao_atual)

    def strip_inline(txt):
        txt = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", txt)
        txt = re.sub(r"`(.+?)`", r"\1", txt)
        return txt

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # bloco de codigo
        if stripped.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                for linha in code_buf:
                    run = p.add_run(linha[:110] + "\n")
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # tabela
        if stripped.startswith("|"):
            flush_table()
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_buf.append(lines[i])
                i += 1
            flush_table()
            continue

        # divisória
        if re.match(r"^---+$", stripped):
            flush_table()
            add_horiz_rule(doc)
            i += 1
            continue

        # H1
        if stripped.startswith("# "):
            flush_table()
            flush_fluxo()
            titulo_h1 = strip_inline(stripped[2:].strip())
            secao_atual = titulo_h1
            doc.add_heading(titulo_h1, level=1)
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            flush_table()
            doc.add_heading(strip_inline(stripped[3:].strip()), level=2)
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            flush_table()
            doc.add_heading(strip_inline(stripped[4:].strip()), level=3)
            i += 1
            continue

        # H4
        if stripped.startswith("#### "):
            flush_table()
            doc.add_heading(strip_inline(stripped[5:].strip()), level=3)
            i += 1
            continue

        # lista
        if re.match(r"^[-*] ", stripped):
            flush_table()
            txt = strip_inline(re.sub(r"^[-*] ", "", stripped))
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(txt)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            i += 1
            continue

        # linha em branco
        if stripped == "" or stripped == "&nbsp;":
            flush_table()
            doc.add_paragraph()
            i += 1
            continue

        # paragrafo normal
        flush_table()
        txt = strip_inline(stripped)
        if txt:
            p = doc.add_paragraph()
            run = p.add_run(txt)
            run.font.name = "Arial"
            run.font.size = Pt(10)
        i += 1

    flush_table()
    flush_fluxo()


# ── Geracao ────────────────────────────────────────────────────────────────────
def gerar(md_path, docx_path, titulo_capa, subtitulo_capa):
    with open(md_path, encoding="utf-8") as f:
        md = f.read()

    partes = re.split(r"(?m)^(?=# )", md, maxsplit=1)
    intro  = partes[0].strip()
    secoes = partes[1] if len(partes) > 1 else ""

    doc = Document()
    setup_styles(doc)

    # Margens A4
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    add_capa(doc, titulo_capa, subtitulo_capa)

    if intro:
        render_md(doc, intro, inserir_fluxo=False)
        page_break(doc)

    if secoes:
        render_md(doc, secoes, inserir_fluxo=True)

    doc.save(docx_path)
    print(f"OK: {docx_path}")


# ── MAIN ───────────────────────────────────────────────────────────────────────
gerar(
    f"{DOCS}/relatorio_gestor.md",
    f"{DOCS}/relatorio_gestor.docx",
    "Painéis Vigilância - Recife",
    "Guia de origem dos dados e visualizações\npara gestores e tomadores de decisão",
)

gerar(
    f"{DOCS}/relatorio_tecnico.md",
    f"{DOCS}/relatorio_tecnico.docx",
    "Painéis Vigilância - Recife",
    "Documentação técnica de dados e variáveis\npor página do painel",
)
