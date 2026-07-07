from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def set_h(para, text, level):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x1D,0x9E,0x75)
    elif level == 2:
        run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x18,0x5F,0xA5)

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    set_h(p, text, 1)

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    set_h(p, text, 2)

def note(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x60,0x60,0x60); r.font.italic = True

def body(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(10); p.paragraph_format.space_after = Pt(4)

def shade(cell, c='E6F1FB'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),c)
    tcPr.append(shd)

def tbl(rows):
    if not rows or len(rows)<2: return
    hdr, data = rows[0], rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(hdr))
    t.style = 'Table Grid'
    for ci,h in enumerate(hdr):
        cell = t.rows[0].cells[ci]; cell.text = h; shade(cell)
        for r in cell.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
    for ri,row in enumerate(data):
        for ci,v in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(v or '')
            for r in cell.paragraphs[0].runs: r.font.size=Pt(9)
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
r = p.add_run('Documentação do Painel Recife — Violência')
r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(0x1D,0x9E,0x75)
p.paragraph_format.space_after=Pt(4)
p2 = doc.add_paragraph()
r2 = p2.add_run('Documentação técnica de dados, variáveis, filtros e visualizações por página')
r2.font.size=Pt(10); r2.font.color.rgb=RGBColor(0x60,0x60,0x60)
p2.paragraph_format.space_after=Pt(16)

# ── P1 ─────────────────────────────────────────────────────────────────
h1('Página 1 - Mapa da Violência')
h2('O que é?')
body('Mapa interativo de Recife com a localização geográfica das mulheres identificadas nos bancos de dados, sobreposta aos bairros da cidade e às unidades de saúde (CNES). Permite visualizar a distribuição espacial dos casos e identificar concentrações por região.')
h2('De onde vêm os dados?')
tbl([['Dado','O que contém'],['Pontos das mulheres','Coordenadas geográficas + trajetória nos sistemas via linkage'],['Bairros de Recife','Polígonos dos bairros'],['Unidades de saúde','Localização e nome (CNES)']])
h2('1 · SQL / ETL')
note('Arquivos .qs e .RData pré-processados — carregados em global.R e mapa_ui.R.')
tbl([['Coluna','Tabela de origem','Descrição'],
['pontos_viol_real.qs','dados/pontos_viol_real.qs','Mulheres com ≥1 notificação SINAN; coordenada do endereço mais recente via linkage probabilístico (SINAN Viol + SIH + SIM + e-SUS APS)'],
['cnes.RData','dados/cnes.RData','Cadastro de unidades de saúde: cd_cnes_unid_not, NO_FANTASIA, latitude_cnes, longitude_cnes'],
['cnes_join.RData','dados/cnes_join.RData','Join CNES ↔ id_unico: cd_cnes_unid_not, texto_final'],
['bairros.geojson','dados/bairros.geojson','Polígonos dos bairros do Recife (objeto sf)'],
['linha_vida_esus4.qs','dados/linha_vida_esus4.qs','Trajetória nos sistemas (df_linha_vida): texto_final e datas de atendimento']])
h2('2 · Preparação')
note('Executado uma vez no carregamento do módulo.')
tbl([['Transformação','Descrição'],
['mutate(Latitude, Longitude, ano_geo)','Converte para numeric/integer'],
['left_join(cnes_join, by=id_unico)','Enriquece df_linha_vida com cd_cnes_unid_not e texto_final'],
['left_join(df_linha_vida2, by=id_unico)','Adiciona trajetória e texto_final a pontos_viol'],
['mutate(cd_cnes_unid_not = as.numeric(...))','Garante tipo numérico para match com selected_cnes'],
['st_read(dados/bairros.geojson)','Carrega polígonos como objeto sf (mapa)']])
h2('3 · Filtros')
note('filtro_id_pessoa não filtra linhas — apenas define highlight=TRUE na linha correspondente.')
tbl([['Filtro','O que faz','Input','Variável'],
['Raça/cor','Mantém só mulheres da raça selecionada','filtro_raca','ds_raca'],
['Idade','Mantém mulheres dentro da faixa de idade','filtro_idade','nu_idade_anos'],
['Ano','Mantém registros do período selecionado','filtro_ano','ano_geo'],
['ID Pessoa','Destaca a mulher no mapa (highlight=TRUE, não remove linhas)','filtro_id_pessoa','id_pessoa'],
['CNES','Mantém só mulheres atendidas naquela unidade','selected_cnes (clique no mapa)','cd_cnes_unid_not']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Mapa — polígonos dos bairros','Contornos dos bairros de Recife.','mapa (sf)','','Nenhum'],
['Mapa — marcadores CNES','Localização das unidades de saúde cadastradas no CNES.','cnes','','selected_cnes'],
['Mapa — marcadores de pontos','Cada ponto representa uma mulher com ≥1 notificação de violência, posicionada pela coordenada do endereço mais recente.','pontos_viol','','Todos; destaque se highlight=TRUE'],
['Mapa — mapa de calor','Concentração espacial dos casos — áreas de maior densidade sem exibir pontos individuais.','pontos_viol','','Todos; ativado por toggle_heatmap']])

# ── P2 ─────────────────────────────────────────────────────────────────
h1('Página 2 - SINAN Violência')
h2('O que é?')
body('Análises sobre as notificações de violência doméstica, sexual e outras violências interpessoais registradas em Recife (2016–2025).')
h2('De onde vêm os dados?')
body('Sistema: SINAN — Ficha de Violência Doméstica, Sexual e/ou Outras Violências Interpessoais. Cada linha = uma notificação individual.')
h2('1 · SQL / ETL')
note('Query na view tratado_sinan_viol_publica (PostgreSQL). Resultado salvo como sinan_viol.qs → objeto df_sinan. Filtro: sg_sexo = F.')
tbl([['Coluna','Tabela de origem','Descrição'],
['id_pessoa','tratado_sinan_viol_publica','Identificador único da mulher (chave de linkage)'],
['nu_idade_anos','tratado_sinan_viol_publica','Idade em anos'],
['ds_raca','tratado_sinan_viol_publica','Raça/cor declarada'],
['ano','tratado_sinan_viol_publica','Ano da notificação'],
['ds_bairro_res','tratado_sinan_viol_publica','Bairro de residência'],
['autor_sexo','tratado_sinan_viol_publica','Sexo do agressor (código — decodificado em viol_ui.R)'],
['autor_alco','tratado_sinan_viol_publica','Suspeita de uso de álcool pelo agressor'],
['les_autop','tratado_sinan_viol_publica','Lesão autoprovocada (código — decodificado em viol_ui.R)'],
['local_ocor','tratado_sinan_viol_publica','Local de ocorrência (código 2 dígitos — decodificado em viol_ui.R)'],
['out_vezes','tratado_sinan_viol_publica','Ocorrência repetida (código — decodificado em viol_ui.R)'],
['viol_fisic ... viol_outr','tratado_sinan_viol_publica','Flags de tipo de violência (0/1)'],
['ag_forca ... ag_outros','tratado_sinan_viol_publica','Flags de meio de agressão (0/1)'],
['def_trans ... def_espec','tratado_sinan_viol_publica','Flags de deficiência da vítima (0/1)'],
['tran_ment, tran_comp','tratado_sinan_viol_publica','Flags de transtorno mental/comportamental (0/1)'],
['rede_sau ... enc_vara','tratado_sinan_viol_publica','Flags de encaminhamento para serviços (0/1) — combinados em pares em viol_ui.R']])
h2('2 · Preparação')
note('Derivações em global.R e viol_ui.R. Flags filtro_viol_* criadas para cada tipo de violência.')
tbl([['Coluna derivada','Lógica'],
['faixa_etaria_padrao','faixa_etaria_func(nu_idade_anos)'],
['filtro_viol_fisica ... filtro_viol_autoprovocada','Flag binária por tipo de violência (usada nos filtros)'],
['ds_tipo_violencia','Descrição concatenada dos tipos de violência presentes'],
['ds_autor_sexo','Decode: 1→Masculino, 2→Feminino, 3→Ambos, else Ignorado'],
['les_autop (decodificado)','Decode: 1→Sim, 2→Não, 9/NA→Ignorado'],
['local_ocor (decodificado)','Decode código 2 dígitos para texto legível'],
['out_vezes (decodificado)','Decode: 1→Sim, 2→Não, 9→Ignorado'],
['rede_enc_sau ... infan_enc_juv','OR dos dois flags de encaminhamento correspondentes']])
h2('3 · Filtros')
note('filtro_violencias usa filter_at — retém linhas onde qualquer flag selecionada = 1.')
tbl([['Filtro','O que faz','Input','Variável'],
['Faixa etária','Mantém mulheres dentro da faixa de idade','filtro_idade','faixa_etaria_padrao'],
['Raça/cor','Mantém só mulheres da raça selecionada','filtro_raca','ds_raca'],
['Ano','Mantém registros do período — não aplicado em freq_ano_graf','filtro_ano','ano'],
['Tipo de violência','Mantém notificações com o tipo marcado','filtro_violencias','filtro_viol_* (flags)'],
['Lesão autoprovocada','Mantém pelo status selecionado','les_autop_fil','les_autop']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Frequência por ano','Evolução do número de notificações ao longo dos anos.','ano, faixa_etaria_padrao, ds_raca, les_autop, viol_*','ano_graf(df) → tab_1(ano)','Todos exceto filtro_ano'],
['Distribuição por faixa etária','Proporção de notificações em cada faixa de idade. Download .xlsx disponível.','faixa_etaria_padrao, ds_raca, les_autop, ano, viol_*','faixa_etaria_graf(df) → tab_1(faixa_etaria_padrao)','Todos os filtros'],
['Distribuição por raça/cor','Proporção de notificações por raça/cor declarada da vítima. Download .xlsx disponível.','ds_raca, faixa_etaria_padrao, les_autop, ano, viol_*','raca_cor_graf(df) → tab_1(ds_raca)','Todos os filtros'],
['Tabela de categorias SINAN','Tabela interativa cruzando uma categoria com uma variável de estratificação. Download .xlsx disponível.','Objetos enc, proc, rel, viol, agc, defic, transt','tab_cat_sinan(filtered_df, extrato, valor)','Todos os filtros'],
['Sexo do agressor × uso de álcool','Barras empilhadas: sexo do agressor × suspeita de uso de álcool.','ds_autor_sexo, autor_alco','tab_2(ds_autor_sexo, autor_alco) + pct_row=TRUE','Todos os filtros'],
['Tabela cruzada configurável','Tabela de dupla entrada montada livremente pelo usuário. Download .xlsx disponível.','faixa_etaria_padrao, ds_raca, ano, out_vezes, local_ocor','tabela_2(df, var_row, var_col) → pivot_wider + adorn_totals','Ignora filtro_violencias']])

# ── P3 ─────────────────────────────────────────────────────────────────
h1('Página 3 - SINAN Intoxicação Exógena')
h2('O que é?')
body('Análises sobre notificações de intoxicação exógena — exposição a substâncias tóxicas (medicamentos, agrotóxicos, drogas). Dados de 2016–2025.')
h2('De onde vêm os dados?')
body('Sistema: SINAN — Ficha de Intoxicação Exógena. Cada linha = uma notificação individual.')
h2('1 · SQL / ETL')
note('Query na view tratado_sinan_iexo_publica. Resultado salvo como df_iexo.qs → objeto df_iexo. Filtro: sg_sexo = F.')
tbl([['Coluna','Tabela de origem','Descrição'],
['id_pessoa','tratado_sinan_iexo_publica','Identificador único da mulher (chave de linkage)'],
['nu_idade_anos','tratado_sinan_iexo_publica','Idade em anos'],
['ds_raca','tratado_sinan_iexo_publica','Raça/cor declarada'],
['ano','tratado_sinan_iexo_publica','Ano da notificação'],
['ds_circunstan','tratado_sinan_iexo_publica','Circunstância da intoxicação (tentativa de suicídio, acidente, etc.)'],
['ds_agente_tox','tratado_sinan_iexo_publica','Substância causadora da intoxicação'],
['ds_tpatend','tratado_sinan_iexo_publica','Tipo de atendimento (ambulatorial, hospitalar, domiciliar, etc.)'],
['ds_hospital','tratado_sinan_iexo_publica','Hospitalização (Sim / Não / Ignorado)']])
h2('2 · Preparação')
note('Apenas faixa_etaria_padrao derivada via faixa_etaria_func().')
h2('3 · Filtros')
note('freq_ano_graf usa df_iexo sem filtro_ano (série histórica completa).')
tbl([['Filtro','O que faz','Input','Variável'],
['Raça/cor','Mantém só mulheres da raça selecionada','filtro_raca','ds_raca'],
['Faixa etária','Mantém mulheres dentro da faixa de idade','filtro_faixa_etaria','faixa_etaria_padrao'],
['Circunstância','Mantém só a circunstância selecionada','filtro_circuns','ds_circunstan'],
['Ano','Mantém registros do período — não aplicado em freq_ano_graf','filtro_ano','ano']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Frequência por ano','Evolução das notificações de intoxicação ao longo dos anos.','df_iexo (sem filtro_ano)','ano_graf(df) → tab_1(ano)','Série histórica completa'],
['Distribuição por faixa etária','Proporção de notificações em cada faixa de idade. Download .xlsx disponível.','faixa_etaria_padrao','faixa_etaria_graf(df) → tab_1(faixa_etaria_padrao)','Todos os filtros'],
['Distribuição por raça/cor','Proporção de notificações por raça/cor declarada. Download .xlsx disponível.','ds_raca','raca_cor_graf(df) → tab_1(ds_raca)','Todos os filtros'],
['Distribuição por circunstância','Distribuição por motivo da intoxicação, ordenada por frequência. Download .xlsx disponível.','ds_circunstan','tab_1(ds_circunstan) → reorder por n → barras horizontais','Todos os filtros'],
['Distribuição por agente tóxico','Distribuição por substância causadora, ordenada por frequência. Download .xlsx disponível.','ds_agente_tox','tab_1(ds_agente_tox) → reorder por n → barras horizontais','Todos os filtros'],
['Tipo de atendimento × hospitalização','Barras empilhadas: tipo de atendimento × hospitalização (Sim/Não/Ignorado). Download .xlsx disponível.','ds_tpatend, ds_hospital','tab_2(ds_tpatend, ds_hospital) + pct_row=TRUE','Todos os filtros']])

# ── P4 ─────────────────────────────────────────────────────────────────
h1('Página 4 - Hospitalizações')
h2('O que é?')
body('Análises sobre internações hospitalares por causas externas (capítulos XIX e XX da CID-10) registradas em Recife (2016–2025).')
h2('De onde vêm os dados?')
body('Sistema: SIH — Sistema de Informações Hospitalares do SUS. Tabela auxiliar: CID-10.')
h2('1 · SQL / ETL')
note('Query na view tratado_sih_publica. Filtros: sg_sexo = F · ds_bairro_res IS NOT NULL · cd_diag_pri caps XIX, XX e Z. Resultado salvo como df_sih.qs.')
tbl([['Coluna','Tabela de origem','Descrição'],
['id_pessoa','tratado_sih_publica','Identificador único da mulher (chave de linkage)'],
['id_pareamento','tratado_sih_publica','Chave de linkage do registro'],
['dt_internacao, dt_saida','tratado_sih_publica','Datas de entrada e saída da internação'],
['dt_nascimento, ano','tratado_sih_publica','Data de nascimento e ano da internação'],
['dias_internacao','Calculada na query','dt_saida − dt_internacao'],
['nu_idade_anos','tratado_sih_publica','Idade em anos'],
['sg_sexo','tratado_sih_publica','Sexo (sempre F)'],
['ds_raca','tratado_sih_publica','Raça/cor declarada'],
['ds_bairro_res','tratado_sih_publica','Bairro de residência'],
['cd_diag_pri','tratado_sih_publica','Diagnóstico principal (CID-10)'],
['cd_diag_sec','tratado_sih_publica','Diagnósticos secundários'],
['cd_proc_rea','tratado_sih_publica','Procedimento realizado'],
['motiv_saida','tratado_sih_publica','Motivo da saída'],
['carater_int','tratado_sih_publica','Caráter da internação'],
['mod_intern','tratado_sih_publica','Modalidade (eletiva, urgência...)'],
['cd_cnes_unid_not','tratado_sih_publica','Unidade de saúde']])
h2('2 · Preparação')
note('faixa_etaria_padrao derivada via faixa_etaria_func(). dias_internacao calculada na query SQL.')
h2('3 · Filtros')
note('freq_ano_graf usa df_sih sem filtro_ano (série histórica completa).')
tbl([['Filtro','O que faz','Input','Variável'],
['Raça/cor','Mantém só mulheres da raça selecionada','filtro_raca','ds_raca'],
['Faixa etária','Mantém mulheres dentro da faixa de idade','filtro_faixa_etaria','faixa_etaria_padrao'],
['Ano','Mantém registros do período — não aplicado em freq_ano_graf','filtro_ano','ano']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Frequência por ano','Evolução do número de internações por causas externas ao longo dos anos.','ano, faixa_etaria_padrao, ds_raca','ano_graf(df) → tab_1(ano)','Todos exceto filtro_ano'],
['Distribuição por faixa etária','Proporção de internações em cada faixa de idade. Download .xlsx disponível.','faixa_etaria_padrao','faixa_etaria_graf(df) → tab_1(faixa_etaria_padrao)','Todos os filtros'],
['Distribuição por raça/cor','Proporção de internações por raça/cor declarada. Download .xlsx disponível.','ds_raca','raca_cor_graf(df) → tab_1(ds_raca)','Todos os filtros'],
['Internações por capítulo da CID-10','Proporção de internações por capítulo da CID-10 (XIX e XX). Barras horizontais ordenadas por frequência.','cd_diag_pri, ano + CAPITULO, DESCRICAO_CAT (via join)','read.csv2(cid_10.csv) → left_join por cd_diag_pri = SUBCAT → tab_1(DESCRICAO_CAT)','Apenas filtro_ano']])

# ── P5 ─────────────────────────────────────────────────────────────────
h1('Página 5 - Mortalidade')
h2('O que é?')
body('Análises sobre óbitos por causas externas (capítulos XIX e XX da CID-10) registrados em Recife (2016–2025).')
h2('De onde vêm os dados?')
body('Sistema: SIM — Sistema de Informações sobre Mortalidade. Tabela auxiliar: CID-10.')
h2('1 · SQL / ETL')
note('Query na view tratado_sim_publica. Filtros: sg_sexo = F · causas externas (CID cap. XX). Resultado salvo como df_sim.qs.')
tbl([['Coluna','Tabela de origem','Descrição'],
['id_pessoa','tratado_sim_publica','Identificador único da mulher (chave de linkage)'],
['cd_causabas','tratado_sim_publica','Causa básica do óbito (CID-10)'],
['dt_obito','tratado_sim_publica','Data do óbito'],
['ano','tratado_sim_publica','Ano do óbito'],
['nu_idade_anos','tratado_sim_publica','Idade em anos'],
['sg_sexo','tratado_sim_publica','Sexo (sempre F)'],
['ds_raca','tratado_sim_publica','Raça/cor declarada'],
['ds_bairro_res','tratado_sim_publica','Bairro de residência']])
h2('2 · Preparação')
note('faixa_etaria_padrao derivada via faixa_etaria_func().')
h2('3 · Filtros')
note('freq_ano_graf usa df_sim sem filtro_ano (série histórica completa).')
tbl([['Filtro','O que faz','Input','Variável'],
['Raça/cor','Mantém só mulheres da raça selecionada','filtro_raca','ds_raca'],
['Faixa etária','Mantém mulheres dentro da faixa de idade','filtro_faixa_etaria','faixa_etaria_padrao'],
['Ano','Mantém registros do período — não aplicado em freq_ano_graf','filtro_ano','ano']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Frequência por ano','Evolução do número de óbitos por causas externas ao longo dos anos.','ano, faixa_etaria_padrao, ds_raca','ano_graf(df) → tab_1(ano)','Série histórica completa'],
['Distribuição por faixa etária','Proporção de óbitos em cada faixa de idade. Download .xlsx disponível.','faixa_etaria_padrao','faixa_etaria_graf(df) → tab_1(faixa_etaria_padrao)','Todos os filtros'],
['Distribuição por raça/cor','Proporção de óbitos por raça/cor declarada. Download .xlsx disponível.','ds_raca','raca_cor_graf(df) → tab_1(ds_raca)','Todos os filtros'],
['Óbitos por capítulo da CID-10','Proporção de óbitos por capítulo da CID-10 (XIX e XX). Barras horizontais ordenadas por frequência.','cd_causabas, ano + CAPITULO, DESCRICAO_CAT (via join)','read.csv2(cid_10.csv) → left_join por cd_causabas = SUBCAT → tab_1(DESCRICAO_CAT)','Apenas filtro_ano']])

# ── P6 ─────────────────────────────────────────────────────────────────
h1('Página 6 - Análise do Linkage')
h2('O que é?')
body('Resultados do linkage probabilístico — cruzamento de SINAN Violência, SIH, SIM e e-SUS APS para identificar mulheres presentes em múltiplos sistemas (2019–2025).')
h2('De onde vêm os dados?')
body('Linkage probabilístico entre SINAN Violência, SIH, SIM e e-SUS APS. Objetos: base_linkage_feminino + linkage_compilado + icd_map.')
h2('1 · SQL / ETL')
note('base_linkage_feminino construída via query que cruza pessoa_publica com registro_linkage, agregando flags de presença por sistema.')
tbl([['Coluna','Tabela de origem','Descrição'],
['id_pessoa','pessoa_publica INNER JOIN registro_linkage','Identificador único da mulher'],
['in_sinan_viol','base_linkage_feminino','Flag: presente no SINAN Violências'],
['in_sim','base_linkage_feminino','Flag: presente no SIM (óbitos)'],
['in_sih','base_linkage_feminino','Flag: presente no SIH (internações)'],
['in_esus','base_linkage_feminino','Flag: presente no e-SUS APS'],
['banco','base_linkage','Sistema de origem do registro'],
['cd_causabas','base_linkage','Código CID-10 da causa (óbito ou internação)'],
['FL_SINAN_VIOL, FL_SINAN_IEXO','base_linkage','Flags de presença por tipo de SINAN'],
['rotulo_viol_interpessoal','linkage_compilado','Flag: tem notificação de violência interpessoal'],
['fl_esus, fl_ob_not_externas','linkage_compilado','Flags de e-SUS e óbito por causas externas'],
['raca_cor, faixa_etaria_padrao','linkage_compilado','Dados demográficos agregados'],
['registros, n_pessoas','linkage_compilado','Contagens de registros e mulheres únicas'],
['cd_causabas → causa_resumida','icd_map_ufmg.Rdata','Dicionário CID-10 → descrição resumida UFMG']])
h2('2 · Preparação')
note('df_obitos e df_internacoes recebem join com icd_map para obter causa_resumida.')
tbl([['Objeto','Lógica'],
['linkage_kpi','select() de linkage_compilado: rotulo_viol_interpessoal, fl_esus, raca_cor, faixa_etaria_padrao, fl_ob_not_externas, registros, n_pessoas'],
['df_obitos','base_linkage |> filter(banco==SIM) |> left_join(icd_map, by=cd_causabas)'],
['df_internacoes','base_linkage |> filter(banco==SIH) |> left_join(icd_map, by=cd_causabas)']])
h2('3 · Filtros')
note('Filtros reativos aplicam-se apenas aos KPIs e gráficos demográficos. Bolhas e comparativo SINAN usam dados estáticos.')
tbl([['Filtro','O que faz','Input','Variável'],
['Faixa etária','Restringe KPIs e gráficos demográficos pela faixa de idade','filtro_idade','faixa_etaria_padrao'],
['Raça/cor','Restringe KPIs e gráficos demográficos pela raça','filtro_raca','raca_cor'],
['Presença e-SUS','Inclui ou exclui mulheres com atendimento no e-SUS APS','filtro_esus','fl_esus']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['KPI cards','Total de registros, mulheres únicas, mulheres com notificação de violência e óbitos por agressão entre notificadas.','kpi_filtrada(): sum(registros), sum(n_pessoas), rotulo_viol_interpessoal==1, fl_ob_not_externas==1','','filtro_idade, raça, e-SUS'],
['Distribuição por faixa etária','Comparação da distribuição etária entre mulheres com e sem notificação de violência. Download .xlsx disponível.','kpi_filtrada() por faixa_etaria_padrao e rotulo_viol_interpessoal','faixa_etaria_graf_npessoas() → soma n_pessoas, calcula pct por grupo','filtro_idade, raça, e-SUS'],
['Distribuição por raça/cor','Comparação da distribuição por raça/cor entre mulheres com e sem notificação de violência. Download .xlsx disponível.','kpi_filtrada() por raca_cor e rotulo_viol_interpessoal','raca_cor_graf_npessoas() → soma n_pessoas, calcula pct por grupo','filtro_idade, raça, e-SUS'],
['Causas de óbito por grupo','Proporções de cada causa de óbito entre 3 grupos: com violência, com intoxicação, sem notificação. Período: 2019–2025. Download .xlsx disponível.','df_obitos (banco==SIM) + causa_resumida','tab_1(causa_resumida) por grupo → bind_rows → bubble chart','Nenhum (estático)'],
['Causas de internação por grupo','Mesmo formato para internações hospitalares. Período: 2019–2025. Download .xlsx disponível.','df_internacoes (banco==SIH) + causa_resumida','tab_1(causa_resumida) por grupo → bind_rows → bubble chart','Nenhum (estático)'],
['SINAN Violências vs. Intoxicação Exógena','Número de mulheres exclusivamente no SINAN Violências, exclusivamente no SINAN Intoxicação e em ambos. Período: 2019–2025. Download .xlsx disponível.','linkage_compilado (so_viol, so_iexo, ambos)','pré-computado','Nenhum (estático)']])

# ── P7 ─────────────────────────────────────────────────────────────────
h1('Página 7 - Linha da Vida')
h2('O que é?')
body('Trajetória individual de cada mulher ao longo do tempo, mostrando eventos registrados em múltiplos sistemas. Resultado do linkage entre SINAN, SIH, SIM e e-SUS APS.')
h2('De onde vêm os dados?')
tbl([['Sistema','Símbolo','O que representa'],
['SINAN Violências','X vermelho','Notificação de violência'],
['SIH','triângulo amarelo','Internação hospitalar'],
['SIM','quadrado azul escuro','Óbito'],
['SINAN Intox. Exógena','losango azul claro','Notificação de intoxicação'],
['e-SUS APS','círculo roxo','Atendimento na atenção básica']])
h2('1 · SQL / ETL')
note('Query linha_vida_esus4: registro_linkage INNER JOIN pessoa_publica; filtra mulheres com ≥1 notificação de violência.')
tbl([['Coluna','Tabela de origem','Descrição'],
['id_pessoa, id_pareamento','registro_linkage INNER JOIN pessoa_publica','Identificadores únicos da mulher e do registro'],
['banco','registro_linkage','Sistema de origem do evento'],
['dt_evento_inicio, dt_evento_fim','registro_linkage','Datas de início e fim do evento'],
['dt_comum','Derivada (coalesce)','Data principal do evento — coalesce das datas disponíveis'],
['ds_raca_padronizada','pessoa_publica','Raça/cor padronizada'],
['nu_idade_anos','pessoa_publica','Idade em anos'],
['nome_ultimo_bairro','pessoa_publica','Bairro do último endereço (NA → Sem informação)'],
['fl_sinan_viol, fl_sim, fl_sih, fl_sinan_iexo, fl_esus_aps','Derivadas em linha_vida.R','Flags de presença por sistema (group_by id_pessoa)'],
['so_sinan, so_esus_aps','Derivadas em linha_vida.R','Flags de exclusividade: só SINAN / só e-SUS'],
['fl_les_autop, fl_viol_*','registro_linkage','Flags de lesão autoprovocada e tipo de violência'],
['texto_final','Derivado (HTML)','Resumo textual da trajetória da mulher']])
h2('2 · Preparação')
tbl([['Coluna derivada','Lógica'],
['fl_esus_aps','group_by(id_pessoa) |> mutate(any(banco == e-SUS APS))'],
['so_sinan','group_by |> mutate(all(banco == SINAN_VIOL))'],
['so_esus_aps','group_by |> mutate(all(banco == e-SUS APS))'],
['faixa_etaria_padrao','faixa_etaria_func(idade_minima) — menor idade registrada por pessoa'],
['texto_banco / cor_banco','Labels e cores por sistema de origem']])
h2('3 · Filtros')
note('Filtros aplicados via botão Filtrar. filtro_banco suporta seleção exclusiva via flags so_sinan/so_esus_aps.')
tbl([['Filtro','O que faz','Input','Variável'],
['Raça/cor','Mantém só mulheres da raça selecionada','filtro_raca','ds_raca_padronizada'],
['Faixa etária','Mantém mulheres dentro da faixa de idade mínima','filtro_faixa_etaria','faixa_etaria_padrao'],
['Sistema de origem','Exibe mulheres somente no SINAN, somente no e-SUS ou em ambos','filtro_banco','so_sinan / so_esus_aps / fl_esus_aps'],
['N° notificações','Mantém mulheres com número de notificações no intervalo','filtro_n_notif (slider)','contagem SINAN_VIOL por id_pessoa'],
['Período','Mantém eventos dentro do intervalo de datas','filtro_periodo (dateRange)','dt_evento_inicio'],
['Bairro','Mantém mulheres do bairro selecionado','filtro_bairro','nome_ultimo_bairro'],
['Unidade notificadora','Mantém mulheres atendidas na unidade selecionada','filtro_unidade','cd_cnes_unid_not'],
['ID Pessoa','Filtra diretamente pelo ID da mulher','filtro_id_pessoa','id_pessoa']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Linha da vida','Cada linha horizontal é uma mulher. Os pontos mostram eventos nos diferentes sistemas ao longo do tempo, com cor e forma por sistema. Eixo X: dt_comum; Eixo Y: id_pareamento (ordenado).','df_linha_filtrado()','','Todos os filtros'],
['Painel de detalhes (ao clicar)','Painel lateral exibido ao clicar em um ponto: mostra raça/cor, idade e resumo textual da mulher selecionada.','ds_raca_padronizada, nu_idade_anos, texto_final','','—'],
['Resumo do recorte (KPIs)','Indicadores resumidos: n° de mulheres únicas, total de registros SINAN Violências, máximo / média / mediana de notificações por mulher.','df_linha_filtrado()','n_distinct(id_pessoa), nrow(SINAN_VIOL), max/mean/median(n_notif)','Todos os filtros'],
['Download (.xlsx)','Exporta a trajetória textual de cada mulher no recorte via openxlsx.','distinct(id_pessoa, texto_final)','','Todos os filtros']])

# ── P8 ─────────────────────────────────────────────────────────────────
h1('Página 8 - Subnotificações')
h2('O que é?')
body('Estimativa dos casos de violência prováveis que não foram notificados no SINAN. Dois níveis de análise: por bairro de residência e por UBS de referência.')
h2('De onde vêm os dados?')
tbl([['Sistema','O que contribui'],
['e-SUS APS','Número de mulheres usuárias da atenção básica por bairro/UBS'],
['SINAN Violências','Número de notificações de violência registradas por bairro/UBS']])
h2('1 · SQL / ETL')
note('Dois queries pré-computados combinando dados e-SUS APS, SINAN Violência e o output do modelo estatístico.')
tbl([['Coluna','Tabela de origem','Descrição'],
['neighborhood, pretty_name','dados_modelo_bairro','Identificador e nome legível do bairro'],
['year','dados_modelo_bairro','Ano de referência'],
['resident_esus_users','dados_modelo_bairro','Mulheres usuárias do e-SUS APS residentes no bairro'],
['resident_sinan_notifications','dados_modelo_bairro','Notificações SINAN de violência de residentes'],
['resident_underreported_cases','dados_modelo_bairro','Estimativa de casos subnotificados (modelo)'],
['resident_suspected_cases','dados_modelo_bairro','Total de casos prováveis (notificados + subnotificados)'],
['population','dados_modelo_bairro','População feminina do bairro'],
['subnotification_rate','dados_modelo_bairro','Taxa de subnotificação (x10.000 para exibição)'],
['cnes, unit_name','dados_modelo_ubs','Código CNES e nome da unidade de saúde'],
['esus_users','dados_modelo_ubs','Mulheres usuárias do e-SUS APS na unidade'],
['sinan_notifications','dados_modelo_ubs','Notificações SINAN na unidade'],
['underreported_cases, suspected_cases','dados_modelo_ubs','Estimativas de subnotificação por UBS'],
['subnotification_rate','dados_modelo_ubs','Taxa de subnotificação por UBS (x10.000)']])
h2('2 · Preparação')
note('Nenhuma derivação adicional no R além do join com pretty_name para exibição legível do bairro.')
h2('3 · Filtros')
note('O filtro de UBS depende do bairro selecionado: observeEvent(neighborhood) atualiza os choices de unit_name dinamicamente.')
tbl([['Filtro','O que faz','Input','Variável'],
['Ano','Mantém somente o(s) ano(s) selecionado(s)','year (pickerInput)','year'],
['Bairro','Mantém somente o(s) bairro(s) selecionado(s)','neighborhood (pickerInput)','neighborhood'],
['UBS','Mantém somente a(s) unidade(s) selecionada(s) — choices atualizados por bairro','unit_name (pickerInput)','unit_name']])
h2('4 · Visualizações')
tbl([['Gráfico','O que mostra','Colunas usadas','Pré-processamento','Filtros'],
['Tabela por bairro','Tabela interativa por bairro: usuárias da atenção básica, notificações SINAN, casos subnotificados, casos prováveis, população feminina e taxa por 10k. Download .xlsx disponível.','modelo_bairro: pretty_name, year, resident_esus_users, resident_sinan_notifications, resident_underreported_cases, resident_suspected_cases, population, subnotification_rate','pré-computado','year, neighborhood'],
['Tabela por UBS','Mesmas colunas da tabela de bairros, acrescentando código CNES e nome da unidade. Download .xlsx disponível.','modelo_ubs: pretty_name (via join), year, cnes, unit_name, esus_users, sinan_notifications, underreported_cases, suspected_cases, subnotification_rate','pré-computado','year, neighborhood, unit_name']])

doc.save(r'C:/Users/gabri/painel-shiny-recife/docs/documentacao_widget.docx')
print('OK')
